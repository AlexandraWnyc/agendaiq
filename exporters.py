"""
exporters.py — Database-driven Excel (Part 1) and Word (Part 2) generation
               for OCA Agenda Intelligence v6.

v6 export pipeline ALWAYS hydrates from the latest DB state, including:
  - ai_summary_for_appearance
  - watch_points_for_appearance
  - analyst_working_notes
  - reviewer_notes
  - finalized_brief
  - leg_history_summary
  - appearance_notes
  - legislative_notes (from matter)
"""
import re, logging
from pathlib import Path
from datetime import datetime

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils import safe_filename, extract_watch_points

log = logging.getLogger("oca-agent")

# ── Excel column definitions ──────────────────────────────────
# v6+ now includes separate columns for analyst/reviewer notes and finalized brief
# so exported files fully reflect researcher-added content.

EXCEL_COLUMNS = [
    ("Cmte Agenda\nDate",          14),
    ("Cmte\nItem #",               12),
    ("BCC Agenda\nDate",           14),
    ("BCC\nItem #",                10),
    ("File #",                     10),
    ("Short Title",                35),
    ("Full Title",                 55),
    ("Legislative\nNotes",         28),
    ("Appearance\nNotes",          25),
    ("Leg History\n(AI Summary)",  45),
    ("File Type",                  12),
    ("Status",                     16),
    ("Control",                    22),
    ("AI Summary\n(Part 1)",       60),
    ("Watch Points",               45),
    ("Analyst\nWorking Notes",     45),
    ("Reviewer Notes",             40),
    ("Finalized Brief",            55),
    ("Workflow\nStatus",           15),
    ("Assigned To",                18),
    ("Reviewer",                   18),
    ("Due Date",                   12),
    ("Carried\nForward",           12),
    ("Prior History\nExists",      14),
]


def _appearance_to_row(app: dict, matter: dict) -> list:
    """Convert a DB appearance+matter dict to an Excel row value list.
    Hydrates every researcher-visible field.  Order must match EXCEL_COLUMNS."""
    part1 = app.get("ai_summary_for_appearance") or ""
    watch = app.get("watch_points_for_appearance") or ""

    # If watch points are embedded in part1, extract them
    if not watch and part1:
        part1, watch = extract_watch_points(part1)

    carried = "Yes" if app.get("carried_forward_from_prior") else "No"
    prior_exists = "Yes" if app.get("prior_appearance_id") else "No"

    return [
        app.get("meeting_date", ""),
        app.get("committee_item_number", ""),
        matter.get("bcc_agenda_date", app.get("bcc_item_number", "")),
        app.get("bcc_item_number", ""),
        app.get("file_number", ""),
        matter.get("short_title", app.get("appearance_title", "")),
        matter.get("full_title", ""),
        matter.get("legislative_notes", ""),
        app.get("appearance_notes", ""),
        app.get("leg_history_summary", ""),
        matter.get("file_type", ""),
        matter.get("current_status", ""),
        matter.get("control_body", ""),
        part1,
        watch,
        app.get("analyst_working_notes", ""),
        app.get("reviewer_notes", ""),
        app.get("finalized_brief", ""),
        app.get("workflow_status", ""),
        app.get("assigned_to", ""),
        app.get("reviewer", ""),
        app.get("due_date", ""),
        carried,
        prior_exists,
    ]


def export_excel(rows_data: list[dict], output_path: Path,
                 highlight_new: bool = True, append: bool = False):
    """
    Write or append rows to an Excel tracking sheet.
    rows_data: list of dicts, each with keys from _appearance_to_row.
    """
    hfont  = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
    hfill  = PatternFill(start_color="2E75B6", end_color="2E75B6", fill_type="solid")
    halign = Alignment(horizontal="center", vertical="center", wrap_text=True)
    border = Border(left=Side(style="thin"), right=Side(style="thin"),
                    top=Side(style="thin"), bottom=Side(style="thin"))
    calign = Alignment(vertical="top", wrap_text=True)
    cfont  = Font(name="Calibri", size=10)
    new_fill = PatternFill(start_color="FFFBEA", end_color="FFFBEA", fill_type="solid")

    if output_path.exists() and append:
        wb = load_workbook(str(output_path))
        ws = wb.active
        row = ws.max_row + 1
        is_new_file = False
    else:
        wb = Workbook()
        ws = wb.active
        ws.title = "Agenda Tracking"
        for ci, (header, width) in enumerate(EXCEL_COLUMNS, 1):
            cell = ws.cell(row=1, column=ci, value=header)
            cell.font  = hfont
            cell.fill  = hfill
            cell.alignment = halign
            cell.border = border
            ws.column_dimensions[get_column_letter(ci)].width = width
        ws.row_dimensions[1].height = 38
        row = 2
        is_new_file = True

    for vals in rows_data:
        for ci, val in enumerate(vals, 1):
            cell = ws.cell(row=row, column=ci, value=str(val) if val else "")
            cell.font      = cfont
            cell.alignment = calign
            cell.border    = border
            if highlight_new and (append or not is_new_file):
                cell.fill = new_fill
        row += 1

    ws.freeze_panes = "A2"
    if row > 2:
        ws.auto_filter.ref = (
            f"A1:{get_column_letter(len(EXCEL_COLUMNS))}{row - 1}"
        )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(output_path))
    log.info(f"  Excel saved: {output_path} ({row - 2} rows)")


def _section(doc, label: str, color=(0x2E, 0x75, 0xB6)):
    """Write a coloured bold section label."""
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(*color)
    return p


def export_word(items_data: list[dict], output_path: Path,
                committee_name: str, meeting_date: str, append: bool = False):
    """Write a Word research brief that includes *every* notes field from the DB."""
    if output_path.exists() and append:
        doc = Document(str(output_path))
    else:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        t = doc.add_heading(committee_name, level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sub.add_run("Research Intelligence Report")
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = meta.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        doc.add_paragraph("")

    doc.add_heading(f"Agenda: {meeting_date}", level=1)

    for item in items_data:
        cn    = item.get("committee_item_number", "")
        mid   = item.get("file_number", "")
        title = item.get("short_title") or item.get("appearance_title", "")

        header = f"Item {cn}" if cn else f"File {mid}"
        if mid and cn:
            header += f"  ·  File #{mid}"
        doc.add_heading(header, level=2)

        if title:
            p = doc.add_paragraph()
            r = p.add_run(title)
            r.font.italic = True
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # ── Workflow chip
        ws = item.get("workflow_status") or "New"
        at = item.get("assigned_to") or ""
        rv = item.get("reviewer") or ""
        dd = item.get("due_date") or ""
        meta_bits = [f"Status: {ws}"]
        if at: meta_bits.append(f"Assigned: {at}")
        if rv: meta_bits.append(f"Reviewer: {rv}")
        if dd: meta_bits.append(f"Due: {dd}")
        p = doc.add_paragraph()
        r = p.add_run("  ·  ".join(meta_bits))
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        # ── Prior history (continuity)
        prior_history = item.get("prior_history") or []
        if prior_history:
            _section(doc, "PRIOR HISTORY")
            for ph in prior_history:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                r = p.add_run(
                    f"\u00B7  {ph.get('meeting_date','')} — {ph.get('body_name','')} "
                    f"[{ph.get('workflow_status','New')}]"
                    + ("  (notes carried forward)" if ph.get("carried_forward_from_prior") else "")
                )
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
            doc.add_paragraph("")

        # ── Agenda Debrief (the narrative body of the Part 1 deliverable)
        part1 = (item.get("ai_summary_for_appearance") or "").strip()
        if part1:
            _section(doc, "AGENDA DEBRIEF")
            _write_clean_text_to_doc(doc, part1)

        # ── Watch points
        wp = (item.get("watch_points_for_appearance") or "").strip()
        if wp:
            _section(doc, "WATCH POINTS", color=(0xB4, 0x53, 0x09))
            _write_clean_text_to_doc(doc, wp)

        # ── Legislative history (so the deliverable stands alone)
        lh = (item.get("leg_history_summary") or "").strip()
        if lh:
            _section(doc, "LEGISLATIVE HISTORY", color=(0x44, 0x44, 0x88))
            _write_clean_text_to_doc(doc, lh)

        # ── Research Notes (timestamped, across stages)
        # Aggregates analyst + reviewer notes from every appearance of this
        # matter in the current meeting package. Deep Research (finalized_brief)
        # is INTENTIONALLY excluded — that tab is reference-only.
        notes_blocks = []
        def _stamp(val, who, when, label):
            val = (val or "").strip()
            if not val:
                return ""
            meta = []
            if who:  meta.append(who)
            if when: meta.append(when)
            tag = f" ({' · '.join(meta)})" if meta else ""
            return f"{label}{tag}:\n{val}"

        wn   = _stamp(item.get("analyst_working_notes"),
                      item.get("analyst_notes_updated_by"),
                      item.get("analyst_notes_updated_at"),
                      "Analyst")
        rn   = _stamp(item.get("reviewer_notes"),
                      item.get("reviewer_notes_updated_by"),
                      item.get("reviewer_notes_updated_at"),
                      "Reviewer")
        for blk in (wn, rn):
            if blk:
                notes_blocks.append(blk)

        if notes_blocks:
            _section(doc, "RESEARCH NOTES", color=(0x05, 0x6F, 0x3A))
            for blk in notes_blocks:
                _write_clean_text_to_doc(doc, blk)
                doc.add_paragraph("")

        # NOTE: finalized_brief (Deep Research Notes) is deliberately NOT
        # exported — it is internal reference material only.

        # Separator
        p = doc.add_paragraph()
        r = p.add_run("\u2500" * 60)
        r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        r.font.size = Pt(8)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    log.info(f"  Word saved: {output_path}")


def _write_clean_text_to_doc(doc, text: str):
    """Parse cleaned text into properly formatted Word paragraphs."""
    header_patterns = re.compile(
        r'^(ITEM\s|Sponsor:|Summary:|District|'
        r'Purpose and Background:|Fiscal Impact:|Additional Notes:|'
        r'RESEARCH CONTEXT:|WATCH POINTS:|PART [12])',
        re.I
    )
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if header_patterns.match(line):
            p = doc.add_paragraph()
            if ":" in line:
                label, _, value = line.partition(":")
                rh = p.add_run(label + ":")
                rh.font.bold = True
                rh.font.size = Pt(11)
                if value.strip():
                    rv = p.add_run("  " + value.strip())
                    rv.font.size = Pt(11)
            else:
                rh = p.add_run(line)
                rh.font.bold = True
                rh.font.size = Pt(11)
        elif line.startswith("- "):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            r = p.add_run("\u00B7  " + bullet_text)
            r.font.size = Pt(11)
        else:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.size = Pt(11)


# ── High-level export orchestrators ──────────────────────────
# These ALWAYS hydrate from the DB at the moment of export, so every
# note the researcher has saved is included in the output.

def export_for_meeting(meeting_id: int, output_dir: Path):
    """Generate Excel + Word for all appearances in a given meeting, hydrating
    fully from the database (so researcher-added notes are never missing)."""
    from repository import get_meeting_by_id, get_appearances_for_meeting, get_matter_by_file_number, get_all_appearances_for_matter

    meeting = get_meeting_by_id(meeting_id)
    if not meeting:
        log.error(f"Meeting {meeting_id} not found")
        return []

    appearances = get_appearances_for_meeting(meeting_id)
    if not appearances:
        log.info(f"No appearances for meeting {meeting_id}")
        return []

    cname = meeting["body_name"]
    mdate = meeting["meeting_date"]
    sn    = safe_filename(cname)

    excel_path = output_dir / f"{sn}_Part1_Tracking.xlsx"
    word_path  = output_dir / f"{sn}_Part2_Research.docx"

    excel_rows  = []
    word_items  = []

    for app in appearances:
        matter = get_matter_by_file_number(app["file_number"])
        if not matter:
            continue

        # Ensure meeting metadata is attached for the row builder
        app["meeting_date"] = mdate

        row = _appearance_to_row(app, matter)
        excel_rows.append(row)

        prior_apps = get_all_appearances_for_matter(matter["id"])
        prior_history = [pa for pa in prior_apps if pa["id"] != app["id"]]

        word_items.append({
            **app,
            "short_title":       matter.get("short_title", ""),
            "legislative_notes": matter.get("legislative_notes", ""),
            "prior_history":     prior_history,
        })

    if excel_rows:
        export_excel(excel_rows, excel_path, highlight_new=False, append=False)
    if word_items:
        export_word(word_items, word_path, cname, mdate, append=False)

    return [p for p in [excel_path, word_path] if p.exists()]


def export_for_appearance(appearance_id: int, output_dir: Path):
    """Export a single appearance to its own Excel + Word files (per-item
    download).  Only the one appearance is included in each output."""
    from repository import (get_appearance_by_id, get_matter_by_file_number,
                            get_meeting_by_id, get_all_appearances_for_matter)

    app = get_appearance_by_id(appearance_id)
    if not app:
        log.error(f"Appearance {appearance_id} not found")
        return []
    matter = get_matter_by_file_number(app["file_number"]) or {}
    meeting = get_meeting_by_id(app["meeting_id"]) or {}
    app["meeting_date"] = meeting.get("meeting_date", "")

    file_num = app.get("file_number") or f"app{appearance_id}"
    sn = safe_filename(f"Item_{file_num}")
    excel_path = output_dir / f"{sn}_Tracking.xlsx"
    word_path  = output_dir / f"{sn}_Research.docx"

    # Excel: single row
    row = _appearance_to_row(app, matter)
    export_excel([row], excel_path, highlight_new=False, append=False)

    # Word: single item with prior history
    prior_apps = get_all_appearances_for_matter(matter["id"]) if matter else []
    prior_history = [pa for pa in prior_apps if pa["id"] != app["id"]]
    word_items = [{
        **app,
        "short_title":       matter.get("short_title", ""),
        "legislative_notes": matter.get("legislative_notes", ""),
        "prior_history":     prior_history,
    }]
    export_word(
        word_items, word_path,
        meeting.get("body_name", ""),
        meeting.get("meeting_date", ""),
        append=False,
    )
    return [p for p in [excel_path, word_path] if p.exists()]


def export_for_date(meeting_date: str, output_dir: Path):
    """Export all committees meeting on a given date."""
    from repository import get_meetings_by_date
    files = []
    meetings = get_meetings_by_date(meeting_date)
    for m in meetings:
        files.extend(export_for_meeting(m["id"], output_dir))
    return files


def export_for_file_number(file_number: str, output_dir: Path):
    """Export all appearances of one file number across all meetings."""
    from repository import get_matter_by_file_number, get_all_appearances_for_matter

    matter = get_matter_by_file_number(file_number)
    if not matter:
        log.error(f"File number {file_number} not found in database")
        return []

    appearances = get_all_appearances_for_matter(matter["id"])
    if not appearances:
        log.info(f"No appearances for file {file_number}")
        return []

    sn = safe_filename(f"File_{file_number}")
    excel_path = output_dir / f"{sn}_History.xlsx"
    word_path  = output_dir / f"{sn}_Research.docx"

    excel_rows = []
    word_items = []
    for app in appearances:
        row = _appearance_to_row(app, matter)
        excel_rows.append(row)
        prior_history = [pa for pa in appearances if pa["id"] < app["id"]]
        word_items.append({
            **app,
            "short_title":       matter.get("short_title", ""),
            "legislative_notes": matter.get("legislative_notes", ""),
            "prior_history":     prior_history,
        })

    if excel_rows:
        export_excel(excel_rows, excel_path, highlight_new=False, append=False)
    body_name = appearances[-1].get("body_name", "Multiple Committees")
    if word_items:
        export_word(
            word_items, word_path,
            f"File {file_number}: {matter.get('short_title','')}",
            "Full History", append=False,
        )

    return [p for p in [excel_path, word_path] if p.exists()]
